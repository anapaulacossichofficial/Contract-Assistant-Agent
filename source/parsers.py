from io import BytesIO
from pathlib import Path
from datetime import datetime
import re
from docx import Document
from pypdf import PdfReader

def extract_text_from_docx(uploaded_file) -> str:
    uploaded_file.seek(0)
    doc = Document(BytesIO(uploaded_file.read()))
    parts = [p.text for p in doc.paragraphs if p.text.strip()]

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                text = cell.text.strip()
                if text:
                    parts.append(text)

    return "\n".join(parts)

def extract_text_from_pdf(uploaded_file) -> str:
    uploaded_file.seek(0)
    reader = PdfReader(BytesIO(uploaded_file.read()))
    parts = []

    for page in reader.pages:
        text = page.extract_text() or ""
        if text.strip():
            parts.append(text)

    return "\n".join(parts)

def extract_text(uploaded_file) -> str:
    suffix = Path(uploaded_file.name.lower()).suffix

    if suffix == ".docx":
        return extract_text_from_docx(uploaded_file)
    if suffix == ".pdf":
        return extract_text_from_pdf(uploaded_file)

    raise ValueError("Formato não suportado. Use .pdf ou .docx")

def normalize_text(text: str) -> str:
    return re.sub(r"\s+", " ", text.lower()).strip()

def parse_money(text: str) -> float:
    patterns = [
        r"valor estimado\s*r\$\s*([\d\.]+,\d{2})",
        r"r\$\s*([\d\.]+,\d{2})",
    ]

    for pattern in patterns:
        match = re.search(pattern, text, re.I)
        if match:
            return float(match.group(1).replace(".", "").replace(",", "."))

    return 0.0

def parse_dates(text: str):
    match = re.search(
        r"inicio\s+em\s+(\d{2}/\d{2}/\d{4}).*?termino\s+em\s+(\d{2}/\d{2}/\d{4})",
        text,
        re.I,
    )

    if not match:
        return None, None

    start_date = datetime.strptime(match.group(1), "%d/%m/%Y")
    end_date = datetime.strptime(match.group(2), "%d/%m/%Y")
    return start_date, end_date